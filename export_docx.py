"""Export one research chat to a .docx a policy researcher can use.

A "chat" is a question the user asked plus the agent's grounded answer and
the bills it cited. The frontend posts {title, question, answer, bill_ids}
to POST /api/export; this module renders a clean Word document.

Design choices:
  * Formatting is deliberately plain — Times New Roman 12pt body, bold
    paragraph text for headers (no Word heading styles, which carry their
    own fonts/colours). It reads like a normal memo, which is what a
    policy researcher pastes into their own work.
  * The agent's answer is Markdown; `_render_markdown` handles the subset
    it emits (ATX headers, bullet / numbered lists, blockquotes, and the
    inline spans **bold**, *italic*, `code`).
  * Cited bills are re-resolved server-side via tool_get_bill so the
    Sources block reflects the authoritative corpus, and each bill-level
    reference gets a derived congress.gov URL. Section-level citations
    ("Sec. 3(b) of H.R. 1736") stay inline — polilabs never fabricates a
    section-level URL.
"""
from __future__ import annotations

import io
import json
import re
from datetime import date, datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

from agent.tools import tool_get_bill

BODY_FONT = "Times New Roman"
GRAY = RGBColor(0x6B, 0x6B, 0x6B)
LINK_BLUE = RGBColor(0x1A, 0x4F, 0x8B)

# ---- congress.gov URL derivation ----

_BILL_TYPE_SLUG = {
    "hr": "house-bill", "s": "senate-bill",
    "hjres": "house-joint-resolution", "sjres": "senate-joint-resolution",
    "hconres": "house-concurrent-resolution", "sconres": "senate-concurrent-resolution",
    "hres": "house-resolution", "sres": "senate-resolution",
}
_BILL_TYPE_LABEL = {
    "hr": "H.R.", "s": "S.", "hjres": "H.J.Res.", "sjres": "S.J.Res.",
    "hconres": "H.Con.Res.", "sconres": "S.Con.Res.", "hres": "H.Res.", "sres": "S.Res.",
}


def _ordinal(n: int) -> str:
    """1 -> '1st', 119 -> '119th' (for the Congress number)."""
    if 11 <= (n or 0) % 100 <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get((n or 0) % 10, "th")
    return f"{n}{suffix}"


def _congress_gov_url(congress: int, bill_type: str, bill_number: int) -> str | None:
    slug = _BILL_TYPE_SLUG.get((bill_type or "").lower())
    if not slug or not congress or not bill_number:
        return None
    return (f"https://www.congress.gov/bill/{_ordinal(congress)}-congress/"
            f"{slug}/{bill_number}")


def _display_bill_number(bill_type: str, bill_number: int) -> str:
    """'hr' 1736 -> 'H.R. 1736'."""
    label = _BILL_TYPE_LABEL.get((bill_type or "").lower(), (bill_type or "").upper())
    return f"{label} {bill_number}".strip()


# ---- low-level paragraph helpers (everything is Times New Roman) ----


def _style_run(run, *, size=12, bold=False, italic=False, color=None, mono=False):
    run.bold = bold
    run.italic = italic
    run.font.name = "Courier New" if mono else BODY_FONT
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return run


def _para(doc, text="", *, size=12, bold=False, italic=False, color=None,
          style=None, space_before=0, space_after=4):
    """A paragraph whose single run is forced to the body font. Used for
    headers (bold) and plain lines alike, so no Word style sneaks in a
    different typeface."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if text:
        _style_run(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p


# ---- Markdown rendering ----

_INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Append `text`, honoring **bold**, *italic*, `code`, in the body font."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _style_run(paragraph.add_run(part[2:-2]), bold=True)
        elif part.startswith("`") and part.endswith("`"):
            _style_run(paragraph.add_run(part[1:-1]), mono=True)
        elif part.startswith("*") and part.endswith("*"):
            _style_run(paragraph.add_run(part[1:-1]), italic=True)
        else:
            _style_run(paragraph.add_run(part))


def _render_markdown(doc: Document, text: str) -> None:
    """Render the agent's Markdown answer block by block. Headers become
    bold paragraphs (not Word heading styles), keeping the whole document
    in one clean typeface."""
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            _para(doc, m.group(2).strip(), bold=True, space_before=8, space_after=2)
            continue

        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            _add_inline_runs(_para(doc, style="List Bullet", space_after=2), m.group(1).strip())
            continue

        m = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if m:
            _add_inline_runs(_para(doc, style="List Number", space_after=2), m.group(1).strip())
            continue

        m = re.match(r"^>\s?(.*)$", line)
        if m:
            p = _para(doc, space_after=4)
            _add_inline_runs(p, m.group(1).strip())
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = GRAY
            p.paragraph_format.left_indent = Pt(18)
            continue

        _add_inline_runs(_para(doc, space_after=6), line.strip())


# ---- Sources block ----


def _resolve_bill(bill_id: str) -> dict[str, Any] | None:
    try:
        bill = json.loads(tool_get_bill(bill_id))
    except (json.JSONDecodeError, TypeError):
        return None
    if not isinstance(bill, dict) or bill.get("error") or bill.get("not_found"):
        return None
    return bill


def _add_sources(doc: Document, bill_ids: list[str]) -> None:
    """One entry per cited bill, re-resolved from the corpus, with a
    congress.gov link a researcher can cite."""
    seen, ids = set(), []
    for b in bill_ids or []:
        if b and b not in seen:
            seen.add(b)
            ids.append(b)
    if not ids:
        return

    _para(doc, "Sources", bold=True, size=12, space_before=12, space_after=4)
    for bill_id in ids:
        bill = _resolve_bill(bill_id)
        p = _para(doc, style="List Bullet", space_after=2)
        if bill is None:
            _style_run(p.add_run(f"{bill_id} (not found in corpus)"), italic=True)
            continue
        congress = bill.get("congress")
        bill_type = bill.get("bill_type", "")
        bill_number = bill.get("bill_number")
        _style_run(p.add_run(f"{_display_bill_number(bill_type, bill_number)}, "
                             f"{_ordinal(congress)} Cong."), bold=True)
        title = bill.get("short_title") or bill.get("title")
        if title:
            _style_run(p.add_run(f" — {title}"))
        if bill.get("sponsor"):
            _style_run(p.add_run(f" (Sponsor: {bill['sponsor']})"))
        url = _congress_gov_url(congress, bill_type, bill_number)
        if url:
            link = _para(doc, style="List Bullet 2", space_after=2)
            _style_run(link.add_run(url), size=10, color=LINK_BLUE)


# ---- top-level builder ----


def build_chat_docx(title: str, question: str, answer: str,
                    bill_ids: list[str]) -> bytes:
    """Render one chat (question + grounded answer + sources) to .docx bytes."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)

    title = (title or question or "Polilabs research export").strip()
    question = (question or "").strip()
    answer = (answer or "").strip()

    _para(doc, title, bold=True, size=14, space_after=2)
    _para(doc, f"Polilabs · generated {datetime.now():%B %-d, %Y}",
          italic=True, size=10, color=GRAY, space_after=10)

    # Show the verbatim question only when the title is something else.
    if question and question != title:
        _para(doc, "Research question", bold=True, space_before=6, space_after=2)
        _para(doc, question, space_after=8)

    if answer:
        _render_markdown(doc, answer)
    else:
        _para(doc, "(no answer recorded)", italic=True)

    _add_sources(doc, bill_ids)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_filename(title: str) -> str:
    """A dated, filesystem-safe filename derived from the chat title."""
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", (title or "").strip()).strip("-").lower()
    slug = (slug[:48] or "research").rstrip("-")
    return f"polilabs-{slug}-{date.today():%Y-%m-%d}.docx"
